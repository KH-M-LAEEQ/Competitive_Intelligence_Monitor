import csv
import io

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from sqlalchemy.orm import Session

from app.models.briefing import Briefing
from app.models.change_log import ChangeLog
from app.models.competitor import Competitor

__all__ = ["export_change_logs_csv", "export_change_logs_docx", "export_briefing_pdf"]


def _workspace_change_logs(db: Session, workspace_id: int) -> list[tuple[ChangeLog, str]]:
    return (
        db.query(ChangeLog, Competitor.name)
        .join(Competitor, ChangeLog.competitor_id == Competitor.id)
        .filter(Competitor.workspace_id == workspace_id)
        .order_by(ChangeLog.created_at.desc())
        .all()
    )


def export_change_logs_csv(db: Session, workspace_id: int) -> bytes:
    rows = _workspace_change_logs(db, workspace_id)

    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow([
        "competitor", "classification", "materiality_score", "created_at", "rationale", "diff"
    ])
    for change_log, competitor_name in rows:
        writer.writerow([
            competitor_name,
            change_log.classification or "",
            change_log.materiality_score if change_log.materiality_score is not None else "",
            change_log.created_at.isoformat() if change_log.created_at else "",
            change_log.rationale or "",
            change_log.diff or "",
        ])

    return buffer.getvalue().encode("utf-8")


def export_change_logs_docx(db: Session, workspace_id: int) -> bytes:
    rows = _workspace_change_logs(db, workspace_id)

    document = Document()
    document.add_heading("Competitive Intelligence — Change Log", level=1)

    if not rows:
        document.add_paragraph("No changes detected yet.")
    else:
        for change_log, competitor_name in rows:
            heading = f"{competitor_name} — {change_log.classification or 'unclassified'}"
            document.add_heading(heading, level=2)

            meta = f"Materiality: {change_log.materiality_score if change_log.materiality_score is not None else 'n/a'}"
            if change_log.created_at:
                meta += f" · Detected: {change_log.created_at.isoformat()}"
            document.add_paragraph(meta)

            if change_log.rationale:
                document.add_paragraph(change_log.rationale)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_briefing_pdf(briefing: Briefing) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER)
    styles = getSampleStyleSheet()

    story = [
        Paragraph(briefing.title, styles["Title"]),
        Spacer(1, 12),
        Paragraph(
            f"Audience: {briefing.audience.value} &middot; "
            f"Digest: {briefing.digest_type.value} &middot; "
            f"Status: {briefing.status.value}",
            styles["Normal"],
        ),
        Spacer(1, 12),
    ]

    for paragraph_text in briefing.body_markdown.split("\n\n"):
        cleaned = paragraph_text.strip().replace("\n", "<br/>")
        if cleaned:
            story.append(Paragraph(cleaned, styles["BodyText"]))
            story.append(Spacer(1, 8))

    doc.build(story)
    return buffer.getvalue()
